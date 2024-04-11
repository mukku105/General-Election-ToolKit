import sys
from docx import Document
from app.models.ac_election_officer import AcElectionOfficer
from app.models.polling_station import PollingStation
from app.models.assembly_const import AssemblyConst
from docx.shared import RGBColor, Pt, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

import roman

from app.models.ps_election_officer import PsElectionOfficer

# Create a new document
doc = Document()

def setup_page(doc, width, height):
    sections = doc.sections
    for section in sections:
        section.page_width = Mm(height)
        section.page_height = Mm(width)
    
def set_margins(doc, top, right, bottom, left):
    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(top)
        section.right_margin = Mm(right)
        section.bottom_margin = Mm(bottom)
        section.left_margin = Mm(left)


def generate_comm_plan(ac_no, file_name):
    global doc
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    style.font.size = Pt(11)
    setup_page(doc, 210, 297)
    set_margins(doc, 5, 15, 8, 15)

    ac_name = AssemblyConst.query.filter_by(ac_no=ac_no).first().ac_name
    ac_text = (ac_no if int(ac_no) > 10 else '0' + ac_no) + ' ' + ac_name.upper()

    add_heading(ac_text, level=2, font_size=18)

    add_officers_table(ac_no)

    save_docs(file_name)

def add_heading(text, level=1, font_size=None):
    heading = doc.add_heading(text, level)
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    heading.paragraph_format.alignment = 1

def get_officers(ac_no):
    officers = PsElectionOfficer.query.join(PollingStation).filter(PollingStation.assembly_const_no==ac_no).all()
    return officers

def add_officers_table(ac_no):
    officer_data = get_officers(ac_no)
    print(officer_data)
    # paragraph_before = doc.add_paragraph()
    # paragraph_before.paragraph_format.space_before = Pt(10)

    # Add a table to the document
    table = doc.add_table(rows=1, cols=6)
    table.autofit = True
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    hdr_cells = table.rows[0].cells
            
    hdr_cells[0].text = 'Sl. No.'
    hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    hdr_cells[0].paragraphs[0].paragraph_format.space_before = Pt(5)
    hdr_cells[0].paragraphs[0].paragraph_format.space_after = Pt(5)

    hdr_cells[1].text = 'No. and name of Polling Station'
    hdr_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    hdr_cells[1].paragraphs[0].paragraph_format.space_before = Pt(3)
    hdr_cells[1].paragraphs[0].paragraph_format.space_after = Pt(3)

    hdr_cells[2].text = 'Name, designation and mobile No. of Presiding officer'
    hdr_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    hdr_cells[2].paragraphs[0].paragraph_format.space_before = Pt(3)
    hdr_cells[2].paragraphs[0].paragraph_format.space_after = Pt(3)

    hdr_cells[3].text = 'Name, designation and mobile No. of Polling officer-1'
    hdr_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    hdr_cells[3].paragraphs[0].paragraph_format.space_before = Pt(3)
    hdr_cells[3].paragraphs[0].paragraph_format.space_after = Pt(3)

    hdr_cells[4].text = 'Name, designation and mobile No. of Micro Observers'
    hdr_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    hdr_cells[4].paragraphs[0].paragraph_format.space_before = Pt(3)
    hdr_cells[4].paragraphs[0].paragraph_format.space_after = Pt(3)

    hdr_cells[5].text = 'Name and mobile No. of BLO'
    hdr_cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    hdr_cells[5].paragraphs[0].paragraph_format.space_before = Pt(3)
    hdr_cells[5].paragraphs[0].paragraph_format.space_after = Pt(3)

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True

    for i, ps in enumerate(officer_data, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_cells[0].paragraphs[0].paragraph_format.space_before = Pt(2)
        row_cells[0].paragraphs[0].paragraph_format.space_after = Pt(2)

        row_cells[1].text = ps.polling_station_code + ' ' + PollingStation.query.filter_by(ps_code=ps.polling_station_code).first().ps_name
        row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_cells[1].paragraphs[0].paragraph_format.space_before = Pt(2)
        row_cells[1].paragraphs[0].paragraph_format.space_after = Pt(2)

        row_cells[2].text = ps.presiding_officer if ps.presiding_officer else ''
        row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_cells[2].paragraphs[0].paragraph_format.space_before = Pt(2)
        row_cells[2].paragraphs[0].paragraph_format.space_after = Pt(2)

        row_cells[3].text = ps.polling_officer_1 if ps.polling_officer_1 else ''
        row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_cells[3].paragraphs[0].paragraph_format.space_before = Pt(2)
        row_cells[3].paragraphs[0].paragraph_format.space_after = Pt(2)

        row_cells[4].text = ps.micro_observers if ps.micro_observers else ''
        row_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_cells[4].paragraphs[0].paragraph_format.space_before = Pt(2)
        row_cells[4].paragraphs[0].paragraph_format.space_after = Pt(2)

        row_cells[5].text = ps.block_level_officer if ps.block_level_officer else ''
        row_cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_cells[5].paragraphs[0].paragraph_format.space_before = Pt(2)
        row_cells[5].paragraphs[0].paragraph_format.space_after = Pt(2)  

    
    # Set the width of the first column
    for row in table.rows:
        row.cells[0].width = Inches(0.6)
        row.cells[3].width = Inches(0.5)

def save_docs(file_name):
    doc.save('app/static/generated_file/comm_plan/' + file_name)